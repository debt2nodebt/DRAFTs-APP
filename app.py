import streamlit as st
from docx import Document
from io import BytesIO
import os
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Define the path to the templates folder
TEMPLATES_DIR = os.path.join(os.path.dirname(__file__), "templates")

# Predefined template paths
Templates = {
    "bank_draft": os.path.join(TEMPLATES_DIR, "Python Bank Draft Template.docx"),
    "bank_draft2": os.path.join(TEMPLATES_DIR, "Python Bank Draft Credit Card.docx"),
    "cessation_draft": os.path.join(TEMPLATES_DIR, "Python Cessation Template.docx"),
    "settlement_draft": os.path.join(TEMPLATES_DIR, "Python_settlement_draft_template.docx")
}

# Function to generate Word draft and return as bytes
def generate_word_draft(template_path, replacements):
    if not os.path.exists(template_path):
        st.error(f"Error: Template file not found at '{template_path}'")
        return None

    doc = Document(template_path)

    # Function to apply Times New Roman font
    def apply_font(run):
        run.font.name = "Times New Roman"
        run.font.size = Pt(11)  # Set font size
        run.font.color.rgb = RGBColor(0, 0, 0)  # Set black color

        # Ensure compatibility with older Word versions
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:ascii"), "Times New Roman")
        rFonts.set(qn("w:hAnsi"), "Times New Roman")
        rPr.append(rFonts)

    # Replace placeholders in paragraphs
    for paragraph in doc.paragraphs:
        full_text = "".join(run.text for run in paragraph.runs)
        for key, value in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, value)

        if paragraph.runs:
            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            apply_font(new_run)

    # Replace placeholders in tables and  apply the real font settings
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                full_text = cell.text
                for key, value in replacements.items():
                    if key in full_text:
                        full_text = full_text.replace(key, value)
                    cell.text = full_text  # Update cell content

                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        apply_font(run)

    # Save document to bytes
    doc_io = BytesIO()
    doc.save(doc_io)
    doc_io.seek(0)
    return doc_io

# Streamlit app layout configuration
st.set_page_config(layout="wide")
st.title("Document Generator App")

# Bank Draft Section
st.subheader("1. Bank Draft for Personal Loan")
with st.form("bank_draft_form"):
    col1, col2 = st.columns(2)
    with col1:
        bank_name = st.text_input("Bank Name", key="bank_draft_bank_name")
        loan_type = st.text_input("Loan Type", key="bank_draft_loan_type")
    with col2:
        loan_number = st.text_input("Loan Number", key="bank_draft_loan_number")
        client_name = st.text_input("Client Name", key="bank_draft_client_name")
        mobile_number = st.text_input("Mobile Number", key="bank_draft_mobile_number")
    
    submitted_bank_draft = st.form_submit_button("Generate Bank Draft")

if submitted_bank_draft and client_name and bank_name:
    replacements = {
        "{BankName}": bank_name,
        "{LoanType}": loan_type,
        "{LoanNumber}": loan_number,
        "{ClientName}": client_name,
        "{MobileNumber}": mobile_number
    }
    
    word_file = generate_word_draft(Templates["bank_draft"], replacements)
    
    if word_file:
        st.download_button(
            label="Download Bank Draft (Word)",
            data=word_file,
            file_name=f"{client_name}_{bank_name}_BankDraft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

st.subheader("2. Bank Draft for Credit Card")
with st.form("bank_draft_formm"):
    col1, col2 = st.columns(2)
    with col1:
        bank_name = st.text_input("Bank Name", key="bank_draft_bank_namee")
        loan_type = st.text_input("Loan Type", key="bank_draft_loan_typee")
    with col2:
        loan_number = st.text_input("Loan Number", key="bank_draft_loan_numberr")
        client_name = st.text_input("Client Name", key="bank_draft_client_namee")
        mobile_number = st.text_input("Mobile Number", key="bank_draft_mobile_numberr")
    
    submitted_bank_draftt = st.form_submit_button("Generate Bank Drafttt")

if submitted_bank_draftt and client_name and bank_name:
    replacements = {
        "{BankName}": bank_name,
        "{LoanType}": loan_type,
        "{LoanNumber}": loan_number,
        "{ClientName}": client_name,
        "{MobileNumber}": mobile_number
    }
    
    word_file = generate_word_draft(Templates["bank_draft2"], replacements)
    
    if word_file:
        st.download_button(
            label="Download Bank Draftt (Word)",
            data=word_file,
            file_name=f"{client_name}_{bank_name}_BankDraft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Settlement Draft Section
st.subheader("3. Settlement Draft")
with st.form("settlement_draft_form"):
    col1, col2, col3 = st.columns(3)
    with col1:
        bank_name = st.text_input("Bank Name", key="settlement_bank_name")
        loan_type = st.text_input("Loan Type", key="settlement_loan_type")
    with col2:
        loan_number = st.text_input("Loan Number", key="settlement_loan_number")
        loan_amount = st.text_input("Loan Amount", key="settlement_loan_amount")
    with col3:
        one_time_payment = st.text_input("One Time Payment", key="settlement_one_time_payment")
        client_name = st.text_input("Client Name", key="settlement_client_name")
        our_mobile_number = st.text_input("Mobile Number", key="settlement_mobile_number")
    
    submitted_settlement_draft = st.form_submit_button("Generate Settlement Draft")

if submitted_settlement_draft and client_name and bank_name:
    replacements = {
        "{BankName}": bank_name,
        "{LoanType}": loan_type,
        "{LoanNumber}": loan_number,
        "{LoanAmount}": loan_amount,
        "{OneTimePayment}": one_time_payment,
        "{ClientName}": client_name,
        "{OurMobileNumber}": our_mobile_number
    }

    word_file = generate_word_draft(Templates["settlement_draft"], replacements)
    
    if word_file:
        st.download_button(
            label="Download Settlement Draft (Word)",
            data=word_file,
            file_name=f"{client_name}_{bank_name}_SettlementDraft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )

# Cessation Draft Section
st.subheader("4. Cessation Draft")
with st.form("cessation_draft_form"):
    col1, col2 = st.columns(2)
    with col1:
        bank_name = st.text_input("Bank Name", key="cessation_bank_name")
        client_name = st.text_input("Client Name", key="cessation_client_name")
    with col2:
        loan_type = st.text_input("Loan Type", key="cessation_loan_type")
        loan_number = st.text_input("Loan Number", key="cessation_loan_number")
        mobile_number = st.text_input("Mobile Number", key="cessation_mobile_number")
    
    submitted_cessation_draft = st.form_submit_button("Generate Cessation Draft")

if submitted_cessation_draft and client_name and bank_name:
    replacements = {
        "{BankName}": bank_name,
        "{ClientName}": client_name,
        "{LoanType}": loan_type,
        "{LoanNumber}": loan_number,
        "{MobileNumber}": mobile_number
    }

    word_file = generate_word_draft(Templates["cessation_draft"], replacements)
    
    if word_file:
        st.download_button(
            label="Download Cessation Draft (Word)",
            data=word_file,
            file_name=f"{client_name}_{bank_name}_CessationDraft.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
